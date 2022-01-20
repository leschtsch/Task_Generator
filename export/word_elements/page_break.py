from docx import Document


class PageBreak:
    """
    Данный класс используется для вставки в docx разрыва страницы.
    """
    def __init__(self):
        pass

    def insert(self, doc: Document):
        """
        Функция, вставляющая разрыв страницы в документ.
        :param doc: python_docx Document, в который нужно вставить разрыв страницы.
        """
        doc.paragraphs[-1].add_run()
        doc.paragraphs[-1].runs[-1].add_break(7)  # код разрыва страницы

import pathlib
import os
from time import time

from docx import Document
from lxml import etree  # чтобы конвертировать MML в OMML
import latex2mathml.converter  # чтобы конвертировать LaTeX в MMl
import pygame

from export.altTeX import alttex
from .picture import Picture


class LatexFormula:
    """
    Данный класс используется для вставки в docx формулы. Некоторые функции теряются при конвертации,
    но в остальном работает.
    """

    def __init__(self, latex: str, space: float = 3.0):
        """
        :param latex: строка с формулой LaTex, без дополнительных символов.
        :param space: отступ строки вокруг формулы.
        """
        self.space = space
        self.latex = latex
        mathml = latex2mathml.converter.convert(self.latex)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(str(next(pathlib.Path(os.getcwd()).glob('**/MML2OMML.xsl'))))
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        self.xml = new_dom.getroot()

    def insert(self, doc: Document):
        """
        Функция, вставляющая формулу в документ.
        :param doc: python_docx Document, в который нужно вставить формулу.
        """
        doc.paragraphs[-1].add_run(' ')
        doc.paragraphs[-1]._element.append(self.xml)
        doc.paragraphs[-1].paragraph_format.line_spacing = self.space


class AlttexFormula:
    def __init__(self, formula: str, fsize: float = 15):
        self.formula = formula
        self.img = alttex.render_math(formula, fsize=fsize)
        self.path = 'temp_img/%.6f.png' % time()
        pygame.image.save(self.img, self.path)
        self.exp_img = Picture(self.path)

    def insert(self, doc: Document):
        self.exp_img.insert(doc)

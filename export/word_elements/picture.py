from docx import Document
from docx.shared import Cm


class Picture:
    """
    Класс, описывающий картинку в документе.
    """
    def __init__(self, path, width=None, height=None, inline=False):
        """
        :param path: путь к картинке.
        :param width: ширина картинки в см.
        :param height: высота картинки в см.
        :param inline: True, если картинку не надо всталять в отдельный параграф.
        Если width и height не заданы - изначальный размер. Если заданы оба - не сохраняет масштаб.
        """
        self.path = path
        self.width = None if width is None else Cm(width)
        self.height = None if height is None else Cm(height)
        self.inline = inline

    def insert(self, doc: Document):
        """
        Функция, вставляющая картинку в документ.
        :param doc: python_docx Document, в который нужно вставить картинку.
        """
        if self.inline:
            doc.paragraphs[-1].runs[-1].add_picture(self.path, width=self.width, height=self.height)
        else:
            doc.add_picture(self.path, width=self.width, height=self.height)

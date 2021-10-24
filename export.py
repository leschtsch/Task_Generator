"""
Данный модуль позволяет сохранять ответы генераторов в файл docx. Для этого надо использовать его функцию export
"""

import re
import os
import typing

from docx import Document, enum  # чтобы работать с вордом
from lxml import etree  # чтобы конвертировать MML в OMML
import latex2mathml.converter  # чтобы конвертировать LaTeX в MMl


def get_path(path):
    """
    Данная функция нужна, чтобы получить имя для сохранения дубликатов файлов.
    Если нужной директории нет, создает ее.
    :param path: путь и имя, куда сохранить файл
    :return: путь и имя файла
    """
    path, name = os.path.split(path)  # получаем директорию и имя файла

    # создание пути, если его нет
    if not os.path.exists(path):
        os.mkdir(path)
    # создание пути, если его нет

    # получение списка файлов в директории с нужным именем и без расширения
    files = os.listdir(path if path else None)  # None, так как os.listdir не принимает ''
    files = list(filter(lambda x: re.match(name + r' ?\d*', x), files))
    files = list(map(lambda x: x.split('.')[0], files))
    files = sorted(files)
    # получение списка файлов в директории с нужным именем и без расширения

    # получение индекса файла при имени
    index = ''
    if not files:
        return os.path.join(path, name)
    for i in files[-1]:
        if i in '1234567890':
            index += i
    index = str(int(index) + 1) if index else '1'
    # получение индекса файла при имени

    name += ' ' + index  # добавляем индекс к имени
    return os.path.join(path, name)  # возврат пути


def convert_latex(latex: str) -> str:
    """
    Данная функция преобразует LaTeX в OMML
    :param latex: строка LaTeX
    :return: строка OMML
    """
    mathml = latex2mathml.converter.convert(latex)  # конвертация LaTeX в MathML для дальнейшей конвертации

    ''' 
    MathML в Office MathML и возврат.
    Я не знаю, как это работает. Вот ссылка, где я это нашел: 
    https://github.com/python-openxml/python-docx/issues/320
    '''
    tree = etree.fromstring(mathml)
    xslt = etree.parse('MML2OMML.xsl')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()
    # MathML в Office MathML и возврат


def export(tasks: typing.List[str], answers: typing.List[str], path: str) -> str:
    """
    Сохраняет занания и ответы в docx и возвращает путь к сохраненному файлу.
    :param tasks: список заданий в виде LaTex. (без $)!
    :param answers: список заданий в виде LaTex. (без $)!
    :param path: путь, куда сохранять файл (с именем, но без расширения)
    :return: путь, куда файл в итоге сохранен
    """
    doc = Document()  # пустой документ

    # вставка заданий
    for i in range(len(tasks)):
        paragraph = doc.add_paragraph()
        paragraph.add_run(str(i + 1) + '.\t')
        paragraph._element.append(convert_latex(tasks[i]))
    # вставка заданий

    # вставка разрыва страницы
    paragraph = doc.add_paragraph()
    paragraph.add_run()
    doc.paragraphs[-1].runs[-1].add_break(enum.text.WD_BREAK.PAGE)
    # вставка разрыва страницы

    # вставка ответов
    for i in range(len(answers)):
        paragraph = doc.add_paragraph()
        paragraph.add_run(str(i + 1) + '.\t')
        paragraph._element.append(convert_latex(answers[i]))
    # вставка ответов

    # сохранение и возврат пути
    path = get_path(path)
    doc.save(f'{get_path(path)}.docx')
    return path
    # сохранение и возврат пути


# проверка работоспособности
if __name__ == '__main__':
    tasks = [r'x^2+2x+1',
             r'x^2-1',
             r'\frac{111}{74}']
    answers = [r'(x+1)^2', r'(x+1)(x-1)', r'1\frac{1}{2}']
    name = 'тест'
    export(tasks, answers, name)
# проверка работоспособности
